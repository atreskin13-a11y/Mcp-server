import sys
import json
import hashlib
import os
from pathlib import Path
from typing import Set, List, Optional, Dict, Any
from vector_db_service import VectorDBWorker
from docx import Document
from fastmcp import FastMCP
from mcp.server.stdio import stdio_server  # можно оставить, даже если явно не используется

# Добавляем путь к tools
sys.path.append(str(Path(__file__).parent))

from tools.worker_file import list_python_files_full_paths, write_file, read_file_txt
from tools.reviewer import collect_project_code

# =========================
# MCP server
# =========================
mcp = FastMCP("file-tools")

# =========================
# RAG settings and helpers
# =========================
BASE_DIR = Path(__file__).parent.resolve()
RAG_DB_DIR = BASE_DIR / "chroma_storage"
RAG_SETTINGS_FILE = BASE_DIR / "rag_settings.json"

DEFAULT_RAG_SETTINGS: Dict[str, Any] = {
    "chunk_size": 1200,
    "chunk_overlap": 150,
    "top_k": 5,
    "include_exts": [".py", ".md", ".txt", ".docx", ".json", ".csv", ".log"],
    "exclude_dirs": [
        "__pycache__",
        ".git",
        "venv",
        "env",
        ".idea",
        ".mypy_cache",
        ".pytest_cache",
        "node_modules",
        ".vscode",
    ],
    "persist_dir": str(RAG_DB_DIR),
    "default_collection_name": "project_rag",
}


# =========================
# Dependency checks
# =========================
def check_required_packages() -> Dict[str, bool]:
    result = {
        "chromadb": False,
        "python-docx": False,
    }

    try:
        import chromadb  # noqa: F401
        result["chromadb"] = True
    except Exception:
        pass

    try:
        from docx import Document as _Document  # noqa: F401
        result["python-docx"] = True
    except Exception:
        pass

    return result


@mcp.tool()
def check_rag_dependencies_tool() -> str:
    """
    Проверяет установку зависимостей для RAG по Word-документам и ChromaDB.
    """
    try:
        result = check_required_packages()
        lines = ["Проверка зависимостей RAG:"]
        lines.append(f"chromadb: {'OK' if result['chromadb'] else 'НЕ УСТАНОВЛЕН'}")
        lines.append(f"python-docx: {'OK' if result['python-docx'] else 'НЕ УСТАНОВЛЕН'}")

        missing = [name for name, ok in result.items() if not ok]
        if missing:
            lines.append("")
            lines.append("Установите недостающие пакеты:")
            if "chromadb" in missing:
                lines.append("pip install chromadb")
            if "python-docx" in missing:
                lines.append("pip install python-docx")
        else:
            lines.append("")
            lines.append("Все необходимые зависимости установлены.")

        return "\n".join(lines)
    except Exception as e:
        return f"Ошибка check_rag_dependencies_tool: {str(e)}"


# =========================
# DOCX reading
# =========================
def read_docx_file(path: Path) -> str:
    """
    Читает .docx-файл и возвращает его текст, включая таблицы.
    """
    try:
        doc = Document(str(path))
        full_text: List[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    full_text.append(" | ".join(cells))

        return "\n".join(full_text).strip()
    except Exception as e:
        return f"Ошибка при чтении файла {path}: {e}"


@mcp.tool()
def read_docx_tool(file_path: str) -> str:
    """
    Читает .docx-файл и возвращает его текст.
    """
    try:
        path = Path(file_path).resolve()
        if not path.exists():
            return f"Файл не найден: {path}"
        if path.suffix.lower() != ".docx":
            return f"Файл не является .docx: {path}"
        return read_docx_file(path)
    except Exception as e:
        return f"Ошибка read_docx_tool: {str(e)}"


# =========================
# Text reading helpers
# =========================
def read_text_file_with_fallbacks(path: Path) -> str:
    for enc in ("utf-8", "cp1251", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="ignore")
        except Exception:
            pass
    return ""


def read_supported_file(path: Path) -> str:
    """
    Читает поддерживаемые типы файлов для RAG.
    """
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return read_docx_file(path)

    if suffix in {".txt", ".md", ".py", ".json", ".csv", ".log"}:
        return read_text_file_with_fallbacks(path)

    return ""


# =========================
# Settings
# =========================
def load_rag_settings() -> Dict[str, Any]:
    settings = DEFAULT_RAG_SETTINGS.copy()
    if RAG_SETTINGS_FILE.exists():
        try:
            data = json.loads(RAG_SETTINGS_FILE.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                settings.update(data)
        except Exception:
            pass
    return settings


def save_rag_settings(settings: Dict[str, Any]) -> None:
    RAG_SETTINGS_FILE.write_text(
        json.dumps(settings, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def normalize_exts(include_exts: Optional[Set[str] | List[str]], fallback_settings: Optional[Dict[str, Any]] = None) -> Set[str]:
    settings = fallback_settings or load_rag_settings()
    values = include_exts if include_exts is not None else settings["include_exts"]
    normalized: Set[str] = set()
    for ext in values:
        ext = str(ext).strip()
        if not ext:
            continue
        if not ext.startswith("."):
            ext = "." + ext
        normalized.add(ext.lower())
    return normalized


def normalize_dirs(exclude_dirs: Optional[Set[str] | List[str]], fallback_settings: Optional[Dict[str, Any]] = None) -> Set[str]:
    settings = fallback_settings or load_rag_settings()
    values = exclude_dirs if exclude_dirs is not None else settings["exclude_dirs"]
    return {str(x).strip() for x in values if str(x).strip()}


@mcp.tool()
def get_rag_settings_tool() -> str:
    """
    Возвращает текущие настройки RAG.
    """
    try:
        settings = load_rag_settings()
        return json.dumps(settings, ensure_ascii=False, indent=2)
    except Exception as e:
        return f"Ошибка get_rag_settings_tool: {str(e)}"


@mcp.tool()
def update_rag_settings_tool(
    chunk_size: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
    top_k: Optional[int] = None,
    include_exts: Optional[Set[str]] = None,
    exclude_dirs: Optional[Set[str]] = None,
    persist_dir: Optional[str] = None,
    default_collection_name: Optional[str] = None,
) -> str:
    """
    Обновляет настройки RAG-поиска и индексации.
    Настройки сохраняются в rag_settings.json рядом с сервером.
    """
    try:
        settings = load_rag_settings()

        if chunk_size is not None:
            if chunk_size <= 0:
                return "Ошибка: chunk_size должен быть > 0"
            settings["chunk_size"] = chunk_size

        if chunk_overlap is not None:
            if chunk_overlap < 0:
                return "Ошибка: chunk_overlap должен быть >= 0"
            settings["chunk_overlap"] = chunk_overlap

        if top_k is not None:
            if top_k <= 0:
                return "Ошибка: top_k должен быть > 0"
            settings["top_k"] = top_k

        if include_exts is not None:
            settings["include_exts"] = sorted(list(normalize_exts(include_exts, settings)))

        if exclude_dirs is not None:
            settings["exclude_dirs"] = sorted(list(normalize_dirs(exclude_dirs, settings)))

        if persist_dir is not None:
            settings["persist_dir"] = str(Path(persist_dir).resolve())

        if default_collection_name is not None and default_collection_name.strip():
            settings["default_collection_name"] = default_collection_name.strip()

        if settings["chunk_overlap"] >= settings["chunk_size"]:
            return "Ошибка: chunk_overlap должен быть меньше chunk_size"

        save_rag_settings(settings)
        return "Настройки RAG успешно обновлены:\n" + json.dumps(settings, ensure_ascii=False, indent=2)
    except Exception as e:
        return f"Ошибка update_rag_settings_tool: {str(e)}"


# =========================
# Chunking and file iteration
# =========================
def split_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size должен быть > 0")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap должен быть >= 0")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap должен быть < chunk_size")

    text = text.strip()
    if not text:
        return []

    chunks: List[str] = []
    start = 0
    n = len(text)

    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = end - chunk_overlap

    return chunks


def iter_project_files(root_dir: str, include_exts: Set[str], exclude_dirs: Set[str]) -> List[Path]:
    root = Path(root_dir).resolve()

    if not root.exists():
        raise FileNotFoundError(f"Директория не найдена: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"Это не директория: {root}")

    result: List[Path] = []

    for current_root, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in exclude_dirs]

        for filename in filenames:
            file_path = Path(current_root) / filename
            if file_path.suffix.lower() in include_exts:
                result.append(file_path)

    return sorted(result)


@mcp.tool()
def preview_rag_files_tool(
    root_dir: str = ".",
    include_exts: Optional[Set[str]] = None,
    exclude_dirs: Optional[Set[str]] = None,
    limit: int = 50,
) -> str:
    """
    Показывает, какие файлы будут выбраны для индексации.
    Удобно для отладки перед build_rag_tool.
    """
    try:
        if limit <= 0:
            return "Ошибка: limit должен быть > 0"

        settings = load_rag_settings()
        final_include_exts = normalize_exts(include_exts, settings)
        final_exclude_dirs = normalize_dirs(exclude_dirs, settings)
        files = iter_project_files(root_dir, final_include_exts, final_exclude_dirs)

        lines = [
            f"Корневая директория: {Path(root_dir).resolve()}",
            f"include_exts={sorted(final_include_exts)}",
            f"exclude_dirs={sorted(final_exclude_dirs)}",
            f"Найдено файлов: {len(files)}",
            "",
        ]

        for path in files[:limit]:
            lines.append(str(path))

        if len(files) > limit:
            lines.append("")
            lines.append(f"Показаны первые {limit} файлов из {len(files)}.")

        return "\n".join(lines)
    except Exception as e:
        return f"Ошибка preview_rag_files_tool: {str(e)}"


# =========================
# Vector DB helpers
# =========================
def get_vector_db_worker(collection_name: Optional[str] = None) -> VectorDBWorker:
    """
    Создает worker для работы с векторной БД через vector_db_service.py
    """
    settings = load_rag_settings()
    persist_dir = settings["persist_dir"]
    Path(persist_dir).mkdir(parents=True, exist_ok=True)

    VectorDBWorker.set_client(str(Path(persist_dir).resolve()))
    final_collection_name = collection_name or settings["default_collection_name"]
    return VectorDBWorker(name_collection=final_collection_name)


def try_clear_collection(worker: VectorDBWorker) -> None:
    """
    Очищает коллекцию, если в ней уже есть записи.
    """
    try:
        existing = worker.collection.get(include=[])
        ids = existing.get("ids", [])
        if ids:
            worker.collection.delete(ids=ids)
    except Exception:
        pass


def make_chunk_id(file_path: Path, chunk_index: int) -> str:
    raw = f"{file_path.resolve()}::chunk::{chunk_index}"
    return hashlib.sha256(raw.encode("utf-8", errors="ignore")).hexdigest()


# =========================
# Basic file tools
# =========================
@mcp.tool()
def write_file_tool(
    file_path: str,
    content: str,
    encoding: str = "utf-8",
    mode: str = "w",
    create_parents: bool = True,
) -> str:
    """
    Записывает контент в файл с автоматическим созданием директорий.
    """
    try:
        write_file(file_path, content, encoding, mode, create_parents)
        return f"Файл {file_path} успешно записан"
    except Exception as e:
        return f"Ошибка записи файла: {str(e)}"


@mcp.tool()
def read_text_file_tool(
    file_path: str,
    encoding: str = "utf-8",
    mode: str = "r",
) -> str:
    """
    Читает контент из файла и возвращает его.
    """
    try:
        result = read_file_txt(file_path, encoding, mode)
        return f"Файл {file_path} успешно прочитан:\n{result}"
    except Exception as e:
        return f"Ошибка чтения файла: {str(e)}"


@mcp.tool()
def collect_project_tool(
    root_dir: str = ".",
    output_file: str = "collected_project_code.py",
    include_exts: Optional[Set[str]] = None,
    exclude_dirs: Optional[Set[str]] = None,
    add_line_numbers: bool = False,
) -> str:
    """
    Рекурсивно обходит дерево проекта и собирает код в один файл.
    Возвращает собранный код.
    """
    include_exts = include_exts if include_exts else {".py", ".md", ".txt"}
    exclude_dirs = exclude_dirs if exclude_dirs else {"__pycache__", ".git", "venv", "env"}
    try:
        collect_project_code(root_dir, output_file, include_exts, exclude_dirs, add_line_numbers)
        result = read_file_txt(output_file)
        return result
    except Exception as e:
        return f"Ошибка: {str(e)}"


@mcp.tool()
def list_python_files_tool(directory: str = ".") -> List[str]:
    """
    Возвращает список абсолютных путей ко всем Python-файлам в директории.
    """
    try:
        files = list_python_files_full_paths(directory)
        return files if files else ["Python-файлы не найдены"]
    except Exception as e:
        return [f"Ошибка: {str(e)}"]


@mcp.tool()
def count_python_files_tool(directory: str = ".") -> List[str]:
    """
    Возвращает количество всех Python-файлов в директории.
    """
    try:
        files = list_python_files_full_paths(directory)
        return [f"Count py-files: {len(files)}"] if files else ["Python-файлы не найдены"]
    except Exception as e:
        return [f"Ошибка: {str(e)}"]


# =========================
# RAG tools
# =========================
@mcp.tool()
def build_rag_tool(
    root_dir: str = ".",
    collection_name: Optional[str] = None,
    include_exts: Optional[Set[str]] = None,
    exclude_dirs: Optional[Set[str]] = None,
    chunk_size: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
    clear_collection: bool = True,
    max_files: int = 500,
    max_file_chars: int = 500000,
    batch_size: int = 50,
) -> str:
    """
    Создает RAG по указанной директории проекта через VectorDBWorker.
    Читает и индексирует файлы по одному.
    """
    try:
        settings = load_rag_settings()

        final_chunk_size = chunk_size if chunk_size is not None else settings["chunk_size"]
        final_chunk_overlap = chunk_overlap if chunk_overlap is not None else settings["chunk_overlap"]
        final_include_exts = normalize_exts(include_exts, settings)
        final_exclude_dirs = normalize_dirs(exclude_dirs, settings)
        final_collection_name = collection_name or settings["default_collection_name"]

        if max_files <= 0:
            return "Ошибка: max_files должен быть > 0"
        if max_file_chars <= 0:
            return "Ошибка: max_file_chars должен быть > 0"
        if batch_size <= 0:
            return "Ошибка: batch_size должен быть > 0"
        if final_chunk_overlap >= final_chunk_size:
            return "Ошибка: chunk_overlap должен быть < chunk_size"

        files = iter_project_files(root_dir, final_include_exts, final_exclude_dirs)
        if not files:
            return (
                "Файлы для индексации не найдены. "
                f"root_dir={Path(root_dir).resolve()}, include_exts={sorted(final_include_exts)}"
            )

        files = files[:max_files]
        worker = get_vector_db_worker(final_collection_name)

        if clear_collection:
            try_clear_collection(worker)

        total_files = 0
        skipped_files = 0
        total_chunks = 0
        errors: List[str] = []
        indexed_files: List[str] = []

        root_resolved = Path(root_dir).resolve()

        for file_path in files:
            try:
                suffix = file_path.suffix.lower()
                text = read_supported_file(file_path)

                if not text:
                    skipped_files += 1
                    continue

                if text.startswith("Ошибка при чтении файла"):
                    skipped_files += 1
                    errors.append(f"{file_path.name}: {text}")
                    continue

                text = text[:max_file_chars].strip()
                if not text:
                    skipped_files += 1
                    continue

                chunks = split_text(text, final_chunk_size, final_chunk_overlap)
                if not chunks:
                    skipped_files += 1
                    continue

                abs_path = str(file_path.resolve())
                rel_path = str(file_path.resolve().relative_to(root_resolved))

                file_ids: List[str] = []
                file_documents: List[str] = []
                file_metadatas: List[Dict[str, Any]] = []

                for i, chunk in enumerate(chunks):
                    file_ids.append(make_chunk_id(file_path, i))
                    file_documents.append(chunk)
                    file_metadatas.append(
                        {
                            "file_path": abs_path,
                            "relative_path": rel_path,
                            "file_name": file_path.name,
                            "suffix": suffix,
                            "chunk_index": i,
                            "source_dir": str(root_resolved),
                        }
                    )

                for start in range(0, len(file_documents), batch_size):
                    end = start + batch_size
                    worker.add_texts(
                        texts=file_documents[start:end],
                        ids=file_ids[start:end],
                        metadatas=file_metadatas[start:end],
                    )

                total_files += 1
                total_chunks += len(file_documents)
                indexed_files.append(rel_path)

            except Exception as e:
                skipped_files += 1
                errors.append(f"{file_path.name}: {e}")
                continue

        if total_chunks == 0:
            if errors:
                return "Не удалось сформировать ни одного чанка для индексации.\n" + "\n".join(errors[:10])
            return "Не удалось сформировать ни одного чанка для индексации."

        result_lines = [
            "RAG успешно создан.",
            f"Коллекция: {final_collection_name}",
            f"Корневая директория: {root_resolved}",
            f"Проиндексировано файлов: {total_files}",
            f"Пропущено файлов: {skipped_files}",
            f"Всего чанков: {total_chunks}",
            f"chunk_size={final_chunk_size}, chunk_overlap={final_chunk_overlap}",
            f"max_files={max_files}, max_file_chars={max_file_chars}, batch_size={batch_size}",
            f"include_exts={sorted(final_include_exts)}",
            f"exclude_dirs={sorted(final_exclude_dirs)}",
            "",
            "Проиндексированные файлы:",
        ]
        result_lines.extend(f"- {name}" for name in indexed_files[:50])

        if len(indexed_files) > 50:
            result_lines.append(f"... и еще {len(indexed_files) - 50} файлов")

        if errors:
            result_lines.append("")
            result_lines.append("Примеры ошибок:")
            result_lines.extend(errors[:10])

        return "\n".join(result_lines)

    except Exception as e:
        return f"Ошибка build_rag_tool: {str(e)}"


@mcp.tool()
def query_rag_tool(
    question: str,
    collection_name: Optional[str] = None,
    top_k: Optional[int] = None,
    include_sources: bool = True,
) -> str:
    """
    Выполняет поиск по ранее собранной RAG-базе и возвращает релевантный контекст.
    """
    try:
        if not question.strip():
            return "Вопрос пустой."

        settings = load_rag_settings()
        final_collection_name = collection_name or settings["default_collection_name"]
        final_top_k = top_k if top_k is not None else settings["top_k"]

        if final_top_k <= 0:
            return "Ошибка: top_k должен быть > 0"

        worker = get_vector_db_worker(final_collection_name)
        result = worker.get_answer(question, n_results=final_top_k)

        documents = result.get("documents", [[]])[0]
        metadatas = result.get("metadatas", [[]])[0]
        distances = result.get("distances", [[]])[0] if result.get("distances") else []

        if not documents:
            return (
                f"По запросу ничего не найдено в коллекции '{final_collection_name}'. "
                "Сначала выполните build_rag_tool."
            )

        blocks: List[str] = []
        for idx, doc in enumerate(documents, start=1):
            meta = metadatas[idx - 1] if idx - 1 < len(metadatas) else {}
            distance_info = ""
            if idx - 1 < len(distances):
                distance_info = f"\nscore/distance: {distances[idx - 1]}"

            if include_sources:
                block = (
                    f"[Фрагмент {idx}]\n"
                    f"Файл: {meta.get('file_path', 'unknown')}\n"
                    f"Относительный путь: {meta.get('relative_path', 'unknown')}\n"
                    f"Чанк: {meta.get('chunk_index', 'unknown')}"
                    f"{distance_info}\n"
                    f"{doc}"
                )
            else:
                block = f"[Фрагмент {idx}]\n{doc}"

            blocks.append(block)

        joined_context = "\n\n" + ("\n" + "-" * 80 + "\n").join(blocks)
        return (
            "Найден релевантный контекст по вопросу:\n"
            f"Вопрос: {question}\n"
            f"Коллекция: {final_collection_name}\n"
            f"Количество фрагментов: {len(documents)}\n"
            f"{joined_context}\n\n"
            "Используй этот контекст для генерации финального ответа."
        )
    except Exception as e:
        return f"Ошибка query_rag_tool: {str(e)}"


@mcp.tool()
def answer_with_rag_tool(
    question: str,
    collection_name: Optional[str] = None,
    top_k: Optional[int] = None,
) -> str:
    """
    Ищет релевантные фрагменты в ChromaDB и формирует краткий ответ
    по найденному контексту.
    """
    try:
        if not question.strip():
            return "Вопрос пустой."

        settings = load_rag_settings()
        final_collection_name = collection_name or settings["default_collection_name"]
        final_top_k = top_k if top_k is not None else settings["top_k"]

        if final_top_k <= 0:
            return "Ошибка: top_k должен быть > 0"

        worker = get_vector_db_worker(final_collection_name)
        result = worker.get_answer(question, n_results=final_top_k)

        documents = result.get("documents", [[]])[0]
        metadatas = result.get("metadatas", [[]])[0]

        if not documents:
            return (
                f"По запросу ничего не найдено в коллекции '{final_collection_name}'. "
                "Сначала выполните build_rag_tool."
            )

        answer_parts: List[str] = []
        sources: List[str] = []

        for i, doc in enumerate(documents):
            cleaned = doc.strip()
            if cleaned:
                answer_parts.append(cleaned)

            meta = metadatas[i] if i < len(metadatas) else {}
            rel_path = meta.get("relative_path", meta.get("file_path", "unknown"))
            chunk_idx = meta.get("chunk_index", "unknown")
            sources.append(f"{rel_path} (chunk {chunk_idx})")

        short_answer = "\n\n".join(answer_parts[:3])
        return (
            "Ответ по найденному контексту:\n"
            f"{short_answer}\n\n"
            "Источники:\n- " + "\n- ".join(sources[:final_top_k])
        )
    except Exception as e:
        return f"Ошибка answer_with_rag_tool: {str(e)}"


@mcp.tool()
def build_rag_from_collected_tool(
    collected_file: str = "collected_project_code.py",
    collection_name: Optional[str] = None,
    chunk_size: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
    clear_collection: bool = True,
) -> str:
    """
    Быстро создает RAG по уже собранному файлу проекта.
    """
    try:
        settings = load_rag_settings()

        final_chunk_size = chunk_size if chunk_size is not None else settings["chunk_size"]
        final_chunk_overlap = chunk_overlap if chunk_overlap is not None else settings["chunk_overlap"]
        final_collection_name = collection_name or settings["default_collection_name"]

        file_path = Path(collected_file).resolve()
        if not file_path.exists():
            return (
                f"Файл для индексации не найден: {file_path}\n"
                "Сначала вызовите collect_project_tool, чтобы собрать проект в один файл."
            )

        text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
        if not text:
            return f"Файл {file_path} пустой, индексация невозможна."

        chunks = split_text(text, final_chunk_size, final_chunk_overlap)
        if not chunks:
            return "Не удалось сформировать чанки из собранного файла."

        worker = get_vector_db_worker(final_collection_name)
        if clear_collection:
            try_clear_collection(worker)

        ids: List[str] = []
        documents: List[str] = []
        metadatas: List[Dict[str, Any]] = []

        for i, chunk in enumerate(chunks):
            ids.append(f"{file_path}::collected_chunk::{i}")
            documents.append(chunk)
            metadatas.append(
                {
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "suffix": file_path.suffix.lower(),
                    "chunk_index": i,
                    "source_type": "collected_project_file",
                }
            )

        worker.add_texts(
            texts=documents,
            ids=ids,
            metadatas=metadatas,
        )

        return (
            "RAG по собранному файлу успешно создан.\n"
            f"Коллекция: {final_collection_name}\n"
            f"Файл: {file_path}\n"
            f"Всего чанков: {len(documents)}\n"
            f"chunk_size={final_chunk_size}, chunk_overlap={final_chunk_overlap}"
        )
    except Exception as e:
        return f"Ошибка build_rag_from_collected_tool: {str(e)}"


@mcp.tool()
def inspect_rag_collection_tool(
    collection_name: Optional[str] = None,
    limit: int = 10,
) -> str:
    """
    Показывает, какие документы реально лежат в коллекции ChromaDB.
    """
    try:
        if limit <= 0:
            return "Ошибка: limit должен быть > 0"

        settings = load_rag_settings()
        final_collection_name = collection_name or settings["default_collection_name"]

        worker = get_vector_db_worker(final_collection_name)
        result = worker.collection.get(limit=limit, include=["documents", "metadatas"])

        ids = result.get("ids", [])
        documents = result.get("documents", [])
        metadatas = result.get("metadatas", [])

        if not ids:
            return f"Коллекция '{final_collection_name}' пуста."

        blocks = []
        for i in range(len(ids)):
            meta = metadatas[i] if i < len(metadatas) else {}
            doc = documents[i] if i < len(documents) else ""
            blocks.append(
                f"[Запись {i + 1}]\n"
                f"id: {ids[i]}\n"
                f"file_path: {meta.get('file_path', 'unknown')}\n"
                f"file_name: {meta.get('file_name', 'unknown')}\n"
                f"chunk_index: {meta.get('chunk_index', 'unknown')}\n"
                f"text:\n{doc}"
            )

        return (
            f"Коллекция: {final_collection_name}\n"
            f"Показано записей: {len(ids)}\n\n"
            + ("\n" + "=" * 80 + "\n").join(blocks)
        )
    except Exception as e:
        return f"Ошибка inspect_rag_collection_tool: {str(e)}"


# =========================
# PDF tools
# =========================
@mcp.tool()
def write_pdf_tool(file_path: str, content: bytes) -> str:
    """
    Записывает бинарный контент в PDF-файл.
    """
    try:
        with open(file_path, "wb") as f:
            f.write(content)
        return f"PDF-файл {file_path} успешно записан"
    except Exception as e:
        return f"Ошибка записи PDF-файла: {str(e)}"


@mcp.tool()
def read_pdf_tool(file_path: str) -> bytes:
    """
    Читает бинарный контент из PDF-файла и возвращает его.
    """
    try:
        with open(file_path, "rb") as f:
            return f.read()
    except Exception as e:
        raise RuntimeError(f"Ошибка чтения PDF-файла: {str(e)}")


@mcp.tool()
def extract_text_from_pdf_tool(file_path: str) -> str:
    """
    Извлекает текст из PDF-файла с помощью библиотеки PyMuPDF (fitz).
    Возвращает извлеченный текст.
    """
    try:
        import fitz

        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except ImportError:
        raise RuntimeError("Библиотека PyMuPDF не установлена. Установите с помощью: pip install PyMuPDF")
    except Exception as e:
        raise RuntimeError(f"Ошибка извлечения текста из PDF: {str(e)}")


@mcp.tool()
def extract_tables_from_pdf_tool(file_path: str) -> List[List[List[str]]]:
    """
    Извлекает таблицы из PDF-файла с помощью библиотеки tabula-py.
    Возвращает список таблиц, где каждая таблица представлена как список строк.
    """
    try:
        import tabula

        tables = tabula.read_pdf(file_path, pages="all", multiple_tables=True)
        return [table.values.tolist() for table in tables]
    except ImportError:
        raise RuntimeError("Библиотека tabula-py не установлена. Установите с помощью: pip install tabula-py")
    except Exception as e:
        raise RuntimeError(f"Ошибка извлечения таблиц из PDF: {str(e)}")


@mcp.tool()
def merge_pdfs_tool(pdf_paths: List[str], output_path: str) -> str:
    """
    Объединяет несколько PDF-файлов в один.
    Возвращает путь к объединенному файлу.
    """
    try:
        from PyPDF2 import PdfMerger

        merger = PdfMerger()
        for path in pdf_paths:
            merger.append(path)
        merger.write(output_path)
        merger.close()
        return f"PDF-файлы успешно объединены в {output_path}"
    except ImportError:
        raise RuntimeError("Библиотека PyPDF2 не установлена. Установите с помощью: pip install PyPDF2")
    except Exception as e:
        raise RuntimeError(f"Ошибка объединения PDF-файлов: {str(e)}")


@mcp.tool()
def split_pdf_tool(file_path: str, output_dir: str) -> List[str]:
    """
    Разделяет PDF-файл на отдельные страницы.
    Возвращает список путей к созданным файлам.
    """
    try:
        from PyPDF2 import PdfReader, PdfWriter

        Path(output_dir).mkdir(parents=True, exist_ok=True)
        reader = PdfReader(file_path)
        output_files: List[str] = []

        for i, page in enumerate(reader.pages):
            writer = PdfWriter()
            writer.add_page(page)
            output_path = f"{output_dir}/page_{i + 1}.pdf"
            with open(output_path, "wb") as f:
                writer.write(f)
            output_files.append(output_path)

        return output_files
    except ImportError:
        raise RuntimeError("Библиотека PyPDF2 не установлена. Установите с помощью: pip install PyPDF2")
    except Exception as e:
        raise RuntimeError(f"Ошибка разделения PDF-файла: {str(e)}")


@mcp.tool()
def encrypt_pdf_tool(file_path: str, output_path: str, user_password: str) -> str:
    """
    Шифрует PDF-файл с помощью пароля.
    Возвращает путь к зашифрованному файлу.
    """
    try:
        from PyPDF2 import PdfReader, PdfWriter

        reader = PdfReader(file_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(user_password=user_password)
        with open(output_path, "wb") as f:
            writer.write(f)
        return f"PDF-файл успешно зашифрован и сохранен в {output_path}"
    except ImportError:
        raise RuntimeError("Библиотека PyPDF2 не установлена. Установите с помощью: pip install PyPDF2")
    except Exception as e:
        raise RuntimeError(f"Ошибка шифрования PDF-файла: {str(e)}")


@mcp.tool()
def decrypt_pdf_tool(file_path: str, output_path: str, user_password: str) -> str:
    """
    Расшифровывает PDF-файл с помощью пароля.
    Возвращает путь к расшифрованному файлу.
    """
    try:
        from PyPDF2 import PdfReader, PdfWriter

        reader = PdfReader(file_path)
        if reader.is_encrypted:
            reader.decrypt(user_password)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)
        return f"PDF-файл успешно расшифрован и сохранен в {output_path}"
    except ImportError:
        raise RuntimeError("Библиотека PyPDF2 не установлена. Установите с помощью: pip install PyPDF2")
    except Exception as e:
        raise RuntimeError(f"Ошибка расшифровки PDF-файла: {str(e)}")


@mcp.tool()
def convert_pdf_to_image_tool(file_path: str, output_dir: str) -> List[str]:
    """
    Конвертирует страницы PDF-файла в изображения (PNG).
    Возвращает список путей к созданным изображениям.
    """
    try:
        import fitz

        Path(output_dir).mkdir(parents=True, exist_ok=True)
        doc = fitz.open(file_path)
        image_files: List[str] = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            pix = page.get_pixmap()
            output_path = f"{output_dir}/page_{page_num + 1}.png"
            pix.save(output_path)
            image_files.append(output_path)
        return image_files
    except ImportError:
        raise RuntimeError("Библиотека PyMuPDF не установлена. Установите с помощью: pip install PyMuPDF")
    except Exception as e:
        raise RuntimeError(f"Ошибка конвертации PDF в изображения: {str(e)}")


if __name__ == "__main__":
    print("MCP server with RAG, DOCX and PDF tools loaded.")
    mcp.run()